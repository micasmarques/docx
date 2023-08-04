from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def set_section_columns(section, num_columns):
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(num_columns))


def add_section_with_columns(doc, num_columns, text):
    doc.add_section()
    section = doc.sections[-1]
    set_section_columns(section, num_columns)
    doc.add_paragraph(text)


def add_section_with_columns_and_margin(doc, num_columns, text):
    doc.add_section()
    section = doc.sections[-1]
    set_section_columns(section, num_columns)
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)


def change_orientation(doc):
    current_section = doc.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_section.orientation = WD_ORIENTATION.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height


def edit_margin(doc):
    secao_antepenultima = doc.sections[-1]
    secao_antepenultima.top_margin = Cm(0.5)
    secao_antepenultima.bottom_margin = Cm(0.5)
    secao_antepenultima.left_margin = Cm(1)
    secao_antepenultima.right_margin = Cm(1)


doc = Document()
header = doc.sections[0].header
htable = header.add_table(1, 2, Inches(6))
htab_cells = htable.rows[0].cells
ht0 = htab_cells[0].add_paragraph()
kh = ht0.add_run()
kh.add_picture('logo.jpg', width=Inches(1))
ht1 = htab_cells[1].add_paragraph('put your header text here')
ht1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
add_page_number(doc.sections[0].footer.paragraphs[0].add_run())
doc.add_heading('GeeksForGeeks', 0)

# Adding paragraph with Increased font size
doc.add_heading('Increased Font Size Paragraph:', 3)
para = doc.add_paragraph().add_run(
    'GeeksforGeeks is a Computer Science portal for geeks.')
# Increasing size of the font
para.font.size = Pt(12)

# Adding paragraph with normal font size
doc.add_heading('Normal Font Size Paragraph:', 3)
doc.add_paragraph(
    'GeeksforGeeks is a Computer Science portal for geeks.')

add_section_with_columns(doc, 2, 'Texto para a primeira coluna. ' * 50)

change_orientation(doc)

add_section_with_columns(doc, 3, 'Texto para as colunas. ' * 150)
add_section_with_columns_and_margin(doc, 3, 'Texto para as colunas. ' * 170)

edit_margin(doc)

doc.save('yourdoc.docx')
